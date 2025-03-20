import streamlit as st
import numpy as np
import spacy
import pandas as pd
import pdfplumber
from docx import Document
from pathlib import Path
from datetime import date, datetime
from io import BytesIO, StringIO
import re
from spacy.matcher import Matcher
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define headers
HEADERS = ['Source', 'Profile received date', 'Department', 'Position', 'Name', 'Gender', 
           'Contact number', 'Mail ID', 'Years of Experience', 'Skills', 'Profile feedback',
           'Mode of Interview', 'Interview Date', 'Interview Time', 'Interview Feedback', 'Last Edited']

class ResumeParser:
    def __init__(self):
        self.nlp = spacy.load('en_core_web_sm')
        self.matcher = Matcher(self.nlp.vocab)
        self.current_date = datetime.now().strftime('%d-%m-%Y')
        self.current_date_and_time = datetime.now().strftime('%d-%m-%Y %I:%M %p')

    def extract_gender_from_resume(self, text: str) -> str:
        male_patterns = [r'\b(Male)\b']
        female_patterns = [r'\b(Female)\b']

        male_regex = re.compile('|'.join(male_patterns), re.IGNORECASE)
        female_regex = re.compile('|'.join(female_patterns), re.IGNORECASE)

        if male_regex.search(text):
            return 'Male'
        elif female_regex.search(text):
            return 'Female'
        return 'Gender not found'

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        with pdfplumber.open(BytesIO(file_content)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    def extract_text_from_docx(self, file_content: bytes) -> str:
        doc = Document(BytesIO(file_content))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    
    def extract_phone_number(self, text: str) -> str:
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?(?:\d{1,4})?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}\s?\d{10}',
            r'\d{10}'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                digits = re.sub(r'\D', '', matches[0])
                if len(digits) >= 10:
                    if len(digits) == 10:
                        return f"+91{digits}"
                    else:
                        return f"+{digits}"
        return ''

    def extract_name(self, text: str) -> str:
        lines = text.split('\n')
        name_patterns = [
            r'^([A-Z][a-z]+ )+[A-Z][a-z]+$',
            r'^([A-Z]+ )+[A-Z]+$',
            r'^[A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+$',
            r'^[A-Z]+\s+[A-Z]\.?\s+[A-Z]+(\s+\([A-Z]+\))?$'
        ]

        for line in lines[:3]:
            line = line.strip()
            for pattern in name_patterns:
                match = re.match(pattern, line)
                if match:
                    return re.sub(r'\s*\([^)]*\)', '', match.group(0)).strip()

        first_lines = ' '.join(lines[:5])
        nlp_text = self.nlp(first_lines)
        pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN', 'OP': '?'}]
        self.matcher.add('NAME', [pattern])
        
        matches = self.matcher(nlp_text)
        for _, start, end in matches:
            span = nlp_text[start:end]
            if len(span.text.split()) >= 2 and not span.text.upper() == span.text:
                return span.text.strip()

        return "Name not found"

    def extract_contact_info(self, text: str) -> dict:
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        email_match = re.search(email_pattern, text)
        
        return {
            'Email': email_match.group() if email_match else None,
            'Phone': self.extract_phone_number(text)
        }

    def extract_department(self, text: str) -> str:
        departments = [
            "electrical", "piping", "mechanical", "process", 
            "civil structural", "instrumentation", "digital", 
            "IT", "accounts", "finance", "HR"
        ]
        
        extracted_departments = []
        for department in departments:
            if re.search(r'\b' + re.escape(department) + r'\b', text, re.IGNORECASE):
                extracted_departments.append(department)
        
        return ', '.join(set(extracted_departments)) if extracted_departments else ""

    def calculate_experience(self, text: str) -> int:
        experience = 0
        year_regex = r"((20|19)(\d{2}))"
        year_range = year_regex + r"[-–]?" + year_regex + r"|(?:present|current|till date|today)"
        date_ranges = re.findall(year_range, text, re.IGNORECASE)

        for date_range in date_ranges:
            start_year = int(date_range[0]) if date_range[0] else None
            end_year = None
            
            if date_range[4].lower() in ['present', 'current', 'till date', 'today']:
                end_year = date.today().year
            elif date_range[3]:
                end_year = int(date_range[3])

            if start_year is not None and end_year is not None:
                experience += end_year - start_year

        numerical_experience_pattern = r'(?i)(\d+)\s*years?(\s*\d+\s*months?)?'
        numerical_experience_matches = re.findall(numerical_experience_pattern, text)

        for match in numerical_experience_matches:
            years = int(match[0])
            months = int(match[1].split()[0]) if match[1] else 0
            total_months = years * 12 + months
            experience += total_months // 12

        return experience

    def extract_skills(self, text: str) -> str:
        skill_patterns = [
            'python', 'java', 'c++', 'javascript', 'html', 'css', 
            'sql', 'machine learning', 'data analysis', 
            'project management', 'agile', 'scrum', 'leadership'
        ]
        
        skills_regex = re.compile(r'\b(' + '|'.join(map(re.escape, skill_patterns)) + r')\b', re.IGNORECASE)
        found_skills = skills_regex.findall(text)
        unique_skills = list(set(found_skills))
        
        return " | ".join(unique_skills) if unique_skills else ""

    def parse_resume(self, file_content: bytes, source: str) -> dict:
        raw_text = ""
        
        if file_content.startswith(b'%PDF'):
            raw_text = self.extract_text_from_pdf(file_content)
        elif file_content.startswith(b'PK'):
            raw_text = self.extract_text_from_docx(file_content)

        name = self.extract_name(raw_text)
        contact_info = self.extract_contact_info(raw_text)
        gender = self.extract_gender_from_resume(raw_text)

        return {
            'Source': source,
            'Name': name,
            'Contact number': contact_info.get('Phone'),
            'Mail ID': contact_info.get('Email'),
            'Gender': gender,
            'Years of Experience': self.calculate_experience(raw_text),
            'Skills': self.extract_skills(raw_text),
            'Department': self.extract_department(raw_text),
            'Profile feedback': ''
        }

def main():
    st.set_page_config(page_title="Resume Parser", layout="wide")
    st.title("Multi-Resume Parser")

    # Initialize session state
    if 'master_data' not in st.session_state:
        st.session_state.master_data = pd.DataFrame(columns=HEADERS)

    # Initialize the parser
    parser = ResumeParser()

    # Sidebar for source selection and file upload
    with st.sidebar:
        st.header("Upload Settings")
        source = st.selectbox(
            "Select source",
            ['Avtar', 'Brunel', 'Career Axes', 'Direct Applicant', 'Fabtech', 'Skapare',
             'Revone8', 'Referral', 'Jobs For Her', 'LinkedIn', 'Naukri', 'Sayantrik',
             'Satyam', 'HCL', 'Tech M', 'Cyient', 'SEPAM', 'Workday / Phenom', 'Walk-in',
             'Indeed', 'PERLG']
        )
        uploaded_files = st.file_uploader(
            "Upload Resumes (PDF/DOCX)",
            type=['pdf', 'docx'],
            accept_multiple_files=True
        )

    # Main content area
    tab1, tab2 = st.tabs(["Parse Resumes", "Master Sheet"])

    # Parse Resumes tab
    with tab1:
        if uploaded_files:
            st.header("Parsed Resumes")
            parsed_data = []
            existing_emails = st.session_state.master_data['Mail ID'].tolist()
            existing_phones = st.session_state.master_data['Contact number'].tolist()
            parsed_entries = set()

            for file in uploaded_files:
                try:
                    bytes_data = file.read()
                    parsed_result = parser.parse_resume(bytes_data, source)
                    parsed_result['Profile received date'] = parser.current_date
                    parsed_result['Last Edited'] = parser.current_date_and_time

                    email = parsed_result.get('Mail ID')
                    phone = parsed_result.get('Contact number')

                    if email in existing_emails or phone in existing_phones or (email, phone) in parsed_entries:
                        parsed_result['Profile feedback'] = 'Duplicate'
                    else:
                        parsed_entries.add((email, phone))

                    parsed_data.append(parsed_result)

                except Exception as e:
                    st.error(f"Error parsing {file.name}: {str(e)}")

            if parsed_data:
                df_parsed = pd.DataFrame(parsed_data)
                st.dataframe(df_parsed, use_container_width=True)

                # Export parsed data
                csv = df_parsed.to_csv(index=False)
                st.download_button(
                    "Download Parsed Resumes",
                    csv,
                    "parsed_resumes.csv",
                    "text/csv",
                    key='download-parsed-csv'
                )

                # Update master sheet button
                if st.button("Update Master Sheet"):
                    st.session_state.master_data = pd.concat(
                        [st.session_state.master_data, df_parsed],
                        ignore_index=True
                    )
                    st.success("Master sheet updated successfully!")

    # Master Sheet tab
    with tab2:
        st.header("Master Sheet")
        
        # Display master sheet
        if not st.session_state.master_data.empty:
            st.dataframe(st.session_state.master_data, use_container_width=True)

            # Export master sheet
            csv = st.session_state.master_data.to_csv(index=False)
            st.download_button(
                "Download Master Sheet",
                csv,
                "master_sheet.csv",
                "text/csv",
                key='download-master-csv'
            )
        else:
            st.info("No data in master sheet yet. Parse some resumes first!")

if __name__ == "__main__":
    main()