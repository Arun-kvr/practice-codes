#brute force:
def largest(arr):
    arr.sort()
    n=len(arr)
    return arr[n-1]
arr=[int(x) for x in input().split()]
print(largest(arr))
#tc=log(n)+n=O(nlogn)

#optimal:
def largest(arr):
    large=arr[0]
    n=len(arr)
    for i in range(0,n-1):
        if arr[i]>large:
            large=arr[i]
    return large
arr=[int(x) for x in input().split()]
print(largest(arr))
#tc=O(n)

===≈===================

import pdfplumber
from docx import Document
import viktor as vkt
from viktor.views import DataGroup, DataItem, DataView, WebResult
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Define the view to display parsed resume data
    @vkt.ActionMethod
    def parse_resume_file(self, params, **kwargs):
        # Get the uploaded file from parameters
        resume_file = params.resume
        
        # Extract text from the resume based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Parse the resume to get the extracted data
        parsed_data = self.parse_resume(resume_text)
        
        # Display parsed data in a tabular view
        data_group = DataGroup("Parsed Resume Data", items=[
            DataItem("Name", parsed_data['Name']),
            DataItem("Email", parsed_data['Email']),
            DataItem("Phone", parsed_data['Phone']),
            DataItem("Education", parsed_data['Education']),
            DataItem("Experience", parsed_data['Experience']),
            DataItem("Skills", parsed_data['Skills']),
        ])
        
        return WebResult(DataView(data=data_group))

××××××××××××××××××××××××××××××××÷÷÷÷÷÷÷

import pdfplumber
from docx import Document
import viktor as vkt
from viktor.views import DataGroup, DataItem, DataView, WebResult
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Define the view to display parsed resume data
    @vkt.controller.Controller.action
    def parse_resume_file(self, params, **kwargs):
        # Get the uploaded file from parameters
        resume_file = params.resume
        
        # Extract text from the resume based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Parse the resume to get the extracted data
        parsed_data = self.parse_resume(resume_text)
        
        # Display parsed data in a tabular view
        data_group = DataGroup(
            "Parsed Resume Data", 
            items=[
                DataItem("Name", parsed_data['Name']),
                DataItem("Email", parsed_data['Email']),
                DataItem("Phone", parsed_data['Phone']),
                DataItem("Education", parsed_data['Education']),
                DataItem("Experience", parsed_data['Experience']),
                DataItem("Skills", parsed_data['Skills']),
            ]
        )
        
        # Return the DataView with a name and the data
        return WebResult(DataView("Resume Parsing Result", data_group))

####################№####################



import pdfplumber
from docx import Document
import viktor as vkt
from viktor.views import DataGroup, DataItem, DataView, WebResult
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):  # Use vkt.Controller directly
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Define the view to display parsed resume data
    @vkt.Controller.action  # Use vkt.Controller.action decorator
    def parse_resume_file(self, params, **kwargs):
        # Get the uploaded file from parameters
        resume_file = params.resume
        
        # Extract text from the resume based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Parse the resume to get the extracted data
        parsed_data = self.parse_resume(resume_text)
        
        # Display parsed data in a tabular view
        data_group = DataGroup(
            "Parsed Resume Data", 
            items=[
                DataItem("Name", parsed_data['Name']),
                DataItem("Email", parsed_data['Email']),
                DataItem("Phone", parsed_data['Phone']),
                DataItem("Education", parsed_data['Education']),
                DataItem("Experience", parsed_data['Experience']),
                DataItem("Skills", parsed_data['Skills']),
            ]
        )
        
        # Return the DataView with a name and the data
        return WebResult(DataView("Resume Parsing Result", data_group))


%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%


import pdfplumber
from docx import Document
import viktor as vkt
from viktor.views import DataGroup, DataItem, DataView, WebResult
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Action to parse the resume file
    @vkt.Action()
    def parse_resume_file(self, params, **kwargs):
        # Get the uploaded file from parameters
        resume_file = params.resume
        
        # Extract text from the resume based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Parse the resume to get the extracted data
        parsed_data = self.parse_resume(resume_text)
        
        # Display parsed data in a tabular view
        data_group = DataGroup(
            "Parsed Resume Data", 
            items=[
                DataItem("Name", parsed_data['Name']),
                DataItem("Email", parsed_data['Email']),
                DataItem("Phone", parsed_data['Phone']),
                DataItem("Education", parsed_data['Education']),
                DataItem("Experience", parsed_data['Experience']),
                DataItem("Skills", parsed_data['Skills']),
            ]
        )
        
        # Return the DataView with a name and the data
        return WebResult(DataView("Resume Parsing Result", data_group))


₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹₹


import pdfplumber
from docx import Document
import viktor as vkt
from viktor.views import DataGroup, DataItem, DataView, WebResult
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Define the view to display parsed resume data
    @vkt.View("Resume Parsing Result", duration_guess=5)
    def parse_resume_view(self, params, **kwargs):
        # Get the uploaded file from parameters
        resume_file = params.resume
        
        # Extract text from the resume based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()
        
        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Parse the resume to get the extracted data
        parsed_data = self.parse_resume(resume_text)
        
        # Display parsed data in a tabular view
        data_group = DataGroup(
            "Parsed Resume Data", 
            items=[
                DataItem("Name", parsed_data['Name']),
                DataItem("Email", parsed_data['Email']),
                DataItem("Phone", parsed_data['Phone']),
                DataItem("Education", parsed_data['Education']),
                DataItem("Experience", parsed_data['Experience']),
                DataItem("Skills", parsed_data['Skills']),
            ]
        )
        
        # Return the DataView with a name and the data
        return DataView(data_group)


@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@


import pdfplumber
from docx import Document
import viktor as vkt
from viktor.parametrization import FileField

class Parametrization(vkt.Parametrization):
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

class Controller(vkt.Controller):
    # Link the parametrization
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Example method that handles the file upload and processes the resume
    def process_resume(self, params, **kwargs):
        resume_file = params.resume

        # Determine the file type and extract text accordingly
        file_ext = resume_file.name.split('.')[-1].lower()

        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

        # Parse the resume text
        parsed_data = self.parse_resume(resume_text)

        # For now, just return the parsed data as a dictionary
        return parsed_data

/////////////////////////////////////////


import pdfplumber
from docx import Document
import viktor as vkt
from viktor.parametrization import FileField, Text

class Parametrization(vkt.Parametrization):
    # File upload field for the resume
    resume = FileField('Upload Resume (PDF/DOCX)', file_types=['.pdf', '.docx'])

    # Text fields to display parsed data after the resume is processed
    name = Text('Name', default='No data yet', visible=False)
    email = Text('Email', default='No data yet', visible=False)
    phone = Text('Phone', default='No data yet', visible=False)
    education = Text('Education', default='No data yet', visible=False)
    experience = Text('Experience', default='No data yet', visible=False)
    skills = Text('Skills', default='No data yet', visible=False)


class Controller(vkt.Controller):
    parametrization = Parametrization

    # Function to extract text from PDF
    def extract_text_from_pdf(self, file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        return text

    # Function to extract text from DOCX
    def extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    # Function to parse the resume and extract key sections
    def parse_resume(self, text):
        sections = {
            'Name': '',
            'Email': '',
            'Phone': '',
            'Education': '',
            'Experience': '',
            'Skills': ''
        }

        lines = text.split('\n')
        for i, line in enumerate(lines):
            if 'email' in line.lower():
                sections['Email'] = line
            elif 'phone' in line.lower() or any(char.isdigit() for char in line):
                sections['Phone'] = line
            elif 'education' in line.lower():
                sections['Education'] = " ".join(lines[i:i+3])
            elif 'experience' in line.lower():
                sections['Experience'] = " ".join(lines[i:i+5])
            elif 'skills' in line.lower():
                sections['Skills'] = " ".join(lines[i:i+3])
            elif i == 0:  # Assume the first line is the name
                sections['Name'] = line

        return sections

    # Method that processes the resume and updates the parametrization
    def process_resume(self, params, **kwargs):
        resume_file = params.resume

        # Extract text from the file based on the file extension
        file_ext = resume_file.name.split('.')[-1].lower()

        if file_ext == 'pdf':
            resume_text = self.extract_text_from_pdf(resume_file.get_value())
        elif file_ext == 'docx':
            resume_text = self.extract_text_from_docx(resume_file.get_value())
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

        # Parse the resume to extract sections
        parsed_data = self.parse_resume(resume_text)

        # Update the fields in the parametrization with the parsed data
        self.parametrization.name.default = parsed_data['Name']
        self.parametrization.email.default = parsed_data['Email']
        self.parametrization.phone.default = parsed_data['Phone']
        self.parametrization.education.default = parsed_data['Education']
        self.parametrization.experience.default = parsed_data['Experience']
        self.parametrization.skills.default = parsed_data['Skills']

        # Make the text fields visible to show the parsed results
        self.parametrization.name.visible = True
        self.parametrization.email.visible = True
        self.parametrization.phone.visible = True
        self.parametrization.education.visible = True
        self.parametrization.experience.visible = True
        self.parametrization.skills.visible = True